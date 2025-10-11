# pip install langchain-google-genai
import os
import re
import pdfplumber
import PyPDF4
from typing import Callable, List, Tuple, Dict
from docx import Document as DocxDocument
import extract_msg
import email
from langchain.docstore.document import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
import time

def read_docx(file_path: str) -> str:
    try:
        doc = DocxDocument(file_path)
        return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    except Exception as e:
        print(f"  - Could not read docx file {os.path.basename(file_path)}: {e}")
        return ""

def read_txt(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read()
    except Exception as e:
        print(f"  - Could not read txt file {os.path.basename(file_path)}: {e}")
        return ""

def read_eml(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            msg = email.message_from_file(file)
            payload = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload += part.get_payload(decode=True).decode('utf-8', errors='ignore')
            else:
                payload = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            return payload
    except Exception as e:
        print(f"  - Could not read eml file {os.path.basename(file_path)}: {e}")
        return ""

def read_msg(file_path: str) -> str:
    try:
        msg = extract_msg.Message(file_path)
        return msg.body
    except Exception as e:
        print(f"  - Could not read msg file {os.path.basename(file_path)}: {e}")
        return ""

def ingest_non_pdf(file_path: str) -> Tuple[List[Tuple[int, str]], Dict[str, str]]:
    _, extension = os.path.splitext(file_path)
    content = ""
    
    if extension == ".docx":
        content = read_docx(file_path)
    elif extension == ".txt":
        content = read_txt(file_path)
    elif extension == ".eml":
        content = read_eml(file_path)
    elif extension == ".msg":
        content = read_msg(file_path)
    else:
        print(f"  - Unsupported file extension: {extension}")
        return [], {}

    return [(1, content)] if content else [], {}

def parse_pdf(file_path: str) -> Tuple[List[Tuple[int, str]], Dict[str, str]]:
    try:
        metadata = extract_metadata_from_pdf(file_path)
        pages = extract_pages_from_pdf(file_path)
        return pages, metadata
    except Exception as e:
        print(f"  - Could not parse PDF {os.path.basename(file_path)}: {e}")
        return [], {}

def extract_metadata_from_pdf(file_path: str) -> dict:
    with open(file_path, "rb") as pdf_file:
        reader = PyPDF4.PdfFileReader(pdf_file, strict=False)
        metadata = reader.getDocumentInfo()
        return {
            "title": metadata.get("/Title", "").strip() if metadata.get("/Title") else "",
            "author": metadata.get("/Author", "").strip() if metadata.get("/Author") else "",
        }

def extract_pages_from_pdf(file_path: str) -> List[Tuple[int, str]]:
    pages = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append((page_num + 1, text))
    except Exception as e:
        print(f"  - Error extracting pages from {os.path.basename(file_path)}: {e}")
    return pages

def merge_hyphenated_words(text: str) -> str:
    return re.sub(r"(\w)-\n(\w)", r"\1\2", text)

def fix_newlines(text: str) -> str:
    return re.sub(r"(?<!\n)\n(?!\n)", " ", text)

def remove_multiple_newlines(text: str) -> str:
    return re.sub(r"\n{2,}", "\n", text)

def clean_text(pages: List[Tuple[int, str]]) -> List[Tuple[int, str]]:
    cleaning_functions = [
        merge_hyphenated_words,
        fix_newlines,
        remove_multiple_newlines,
    ]
    cleaned_pages = []
    for page_num, text in pages:
        for cleaning_function in cleaning_functions:
            text = cleaning_function(text)
        cleaned_pages.append((page_num, text))
    return cleaned_pages

def text_to_docs(text: List[Tuple[int, str]], metadata: Dict[str, str], filename: str) -> List[Document]:
    doc_chunks = []
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=30000,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
        chunk_overlap=3000,
    )
    for page_num, page_content in text:
        chunks = text_splitter.split_text(page_content)
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "page_number": page_num,
                    "chunk": i,
                    "source": f"p{page_num}-{i}",
                    "filename": filename,
                    **metadata,
                },
            )
            doc_chunks.append(doc)
    return doc_chunks

def process_files_in_batches(all_files: List[str], batch_size: int, api_key: str):
    total_files = len(all_files)

    for i in range(0, total_files, batch_size):
        batch_files = all_files[i:i + batch_size]
        print(f"\n--- Processing Batch {i//batch_size + 1}/{(total_files + batch_size - 1)//batch_size} ---")
        
        batch_document_chunks = []
        for file_path in batch_files:
            filename = os.path.basename(file_path)
            print(f"Processing File: {filename}...")
            
            try:
                _, extension = os.path.splitext(file_path)
                if extension.lower() == ".pdf":
                    raw_pages, metadata = parse_pdf(file_path)
                else:
                    raw_pages, metadata = ingest_non_pdf(file_path)
                
                if not raw_pages:
                    print(f"  - No content extracted from {filename}. Skipping.")
                    continue

                cleaned_pages = clean_text(raw_pages)
                document_chunks = text_to_docs(cleaned_pages, metadata, filename)
                
                batch_document_chunks.extend(document_chunks)
                print(f"  - Extracted {len(document_chunks)} text chunks from {filename}.")

            except Exception as e:
                print(f"  - An unexpected error occurred while processing {filename}: {e}")

        if not batch_document_chunks:
            print("No documents were processed in this batch. Moving to the next.")
            continue
            
        print(f"\nGenerating embeddings for {len(batch_document_chunks)} text chunks in this batch.")
        try:
           
            embeddings = GoogleGenerativeAIEmbeddings(
                model="models/embedding-001", 
                google_api_key=api_key
            )
            vector_store = Chroma.from_documents(
                documents=batch_document_chunks,
                embedding=embeddings,
                collection_name="ultima-collection",
                persist_directory="chromadata",
            )
            vector_store.persist()
            print("Embeddings complete for this batch, Chroma DB updated.")
        except Exception as e:
            print(f"  - Failed to generate embeddings for this batch: {e}")

def main():
    print("Ingestion script activated...")

    api_key = "AIzaSyD5TjkKiPBLCEiKt-YJwLOhC2XND3nF_T8" 
    
    if not api_key or api_key == "your_google_api_key":
        print("\nERROR: Google API key is missing.")
        print("Please add your key to the `api_key` variable in the main() function.")
        return

    root_folder_path = r"C:\SynologyDrive\Documents\Sollex\Alex's Law Codex"

    print(f"Searching for documents in '{root_folder_path}' and all subfolders...")
    all_files = []
    supported_extensions = ('.pdf', '.docx', '.txt', '.eml', '.msg')
    for root, _, files in os.walk(root_folder_path):
        for file in files:
            if file.lower().endswith(supported_extensions):
                all_files.append(os.path.join(root, file))
    
    if not all_files:
        print(f"No supported documents found in '{root_folder_path}'. Please check the path and file types.")
        return

    print(f"Found {len(all_files)} total files to process.")
    
    batch_size = 20 

    process_files_in_batches(all_files, batch_size, api_key)

    print("\nIngestion completed!")

if __name__ == "__main__":
    main()